import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import docx
import PyPDF2
import os
import webbrowser
import tempfile
from pathlib import Path
import re
import fitz  # PyMuPDF

class DocumentLinker:
    def __init__(self, root):
        self.root = root
        self.root.title("Document Word Linker")
        self.root.geometry("800x600")
        
        # File paths
        self.word_file_path = None
        self.pdf_file_path = None
        
        # Extracted content
        self.word_first_word = None
        self.pdf_second_page_first_word = None
        self.pdf_second_page_first_word_coords = None
        
        # Setup UI
        self.create_widgets()
    
    def create_widgets(self):
        # Frame for file selection
        file_frame = ttk.LabelFrame(self.root, text="Document Selection")
        file_frame.pack(fill="x", padx=10, pady=10)
        
        # Word document selection
        ttk.Label(file_frame, text="Word Document:").grid(row=0, column=0, padx=5, pady=5, sticky="w")
        self.word_file_label = ttk.Label(file_frame, text="No file selected")
        self.word_file_label.grid(row=0, column=1, padx=5, pady=5, sticky="w")
        ttk.Button(file_frame, text="Browse...", command=self.select_word_file).grid(row=0, column=2, padx=5, pady=5)
        
        # PDF document selection
        ttk.Label(file_frame, text="PDF Document:").grid(row=1, column=0, padx=5, pady=5, sticky="w")
        self.pdf_file_label = ttk.Label(file_frame, text="No file selected")
        self.pdf_file_label.grid(row=1, column=1, padx=5, pady=5, sticky="w")
        ttk.Button(file_frame, text="Browse...", command=self.select_pdf_file).grid(row=1, column=2, padx=5, pady=5)
        
        # Frame for process and analysis
        analysis_frame = ttk.LabelFrame(self.root, text="Document Analysis")
        analysis_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        # Word document first word display
        ttk.Label(analysis_frame, text="First word in Word document:").grid(row=0, column=0, padx=5, pady=5, sticky="w")
        self.word_first_word_label = ttk.Label(analysis_frame, text="-")
        self.word_first_word_label.grid(row=0, column=1, padx=5, pady=5, sticky="w")
        
        # PDF second page first word display
        ttk.Label(analysis_frame, text="First word on second page of PDF:").grid(row=1, column=0, padx=5, pady=5, sticky="w")
        self.pdf_first_word_label = ttk.Label(analysis_frame, text="-")
        self.pdf_first_word_label.grid(row=1, column=1, padx=5, pady=5, sticky="w")
        
        # Buttons
        button_frame = ttk.Frame(self.root)
        button_frame.pack(fill="x", padx=10, pady=10)
        
        ttk.Button(button_frame, text="Process Documents", command=self.process_documents).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Create Linked Viewer", command=self.create_linked_viewer).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Exit", command=self.root.quit).pack(side=tk.RIGHT, padx=5)
        
        # Status bar
        self.status_var = tk.StringVar()
        self.status_var.set("Ready")
        status_bar = ttk.Label(self.root, textvariable=self.status_var, relief=tk.SUNKEN, anchor=tk.W)
        status_bar.pack(side=tk.BOTTOM, fill=tk.X)
    
    def select_word_file(self):
        file_path = filedialog.askopenfilename(
            title="Select Word Document",
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
        )
        if file_path:
            self.word_file_path = file_path
            self.word_file_label.config(text=os.path.basename(file_path))
            self.status_var.set(f"Word document selected: {os.path.basename(file_path)}")
    
    def select_pdf_file(self):
        file_path = filedialog.askopenfilename(
            title="Select PDF Document",
            filetypes=[("PDF Documents", "*.pdf"), ("All Files", "*.*")]
        )
        if file_path:
            self.pdf_file_path = file_path
            self.pdf_file_label.config(text=os.path.basename(file_path))
            self.status_var.set(f"PDF document selected: {os.path.basename(file_path)}")
    
    def extract_word_content(self):
        try:
            doc = docx.Document(self.word_file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            if not paragraphs:
                messagebox.showwarning("Warning", "No text found in Word document")
                return False
            
            # Get the first word from first paragraph
            first_paragraph = paragraphs[0]
            first_word = first_paragraph.split()[0] if first_paragraph.split() else ""
            
            self.word_first_word = first_word
            self.word_first_word_label.config(text=first_word)
            
            return True
        except Exception as e:
            messagebox.showerror("Error", f"Error extracting Word content: {str(e)}")
            return False
    
    def extract_pdf_content(self):
        try:
            # Using PyMuPDF (fitz) to extract text and positions
            pdf_document = fitz.open(self.pdf_file_path)
            
            if pdf_document.page_count < 2:
                messagebox.showwarning("Warning", "PDF has fewer than 2 pages")
                return False
            
            # Get second page
            page = pdf_document[1]  # 0-based index, so 1 is the second page
            page_text = page.get_text()
            
            if not page_text.strip():
                messagebox.showwarning("Warning", "No text found on second page of PDF")
                return False
            
            # Get first word from second page and its coordinates
            words = page.get_text("words")  # List of (x0, y0, x1, y1, word, block_no, line_no, word_no)
            
            if not words:
                messagebox.showwarning("Warning", "No words found on second page of PDF")
                return False
            
            first_word_info = words[0]
            first_word = first_word_info[4]  # The word is at index 4
            word_coords = first_word_info[:4]  # Coordinates are at indices 0-3 (x0,y0,x1,y1)
            
            self.pdf_second_page_first_word = first_word
            self.pdf_second_page_first_word_coords = word_coords
            self.pdf_first_word_label.config(text=first_word)
            
            return True
        except Exception as e:
            messagebox.showerror("Error", f"Error extracting PDF content: {str(e)}")
            return False
    
    def process_documents(self):
        if not self.word_file_path or not self.pdf_file_path:
            messagebox.showwarning("Warning", "Please select both Word and PDF documents")
            return
        
        self.status_var.set("Processing documents...")
        self.root.update_idletasks()
        
        word_success = self.extract_word_content()
        pdf_success = self.extract_pdf_content()
        
        if word_success and pdf_success:
            self.status_var.set(f"Linking '{self.word_first_word}' to '{self.pdf_second_page_first_word}'")
            messagebox.showinfo("Success", "Documents processed successfully. Ready to create linked viewer.")
        else:
            self.status_var.set("Processing failed")
    
    def create_linked_viewer(self):
        if not self.word_first_word or not self.pdf_second_page_first_word:
            messagebox.showwarning("Warning", "Please process documents first")
            return
        
        try:
            # Create a PDF.js-based HTML viewer
            temp_dir = tempfile.gettempdir()
            html_path = os.path.join(temp_dir, "document_link_viewer.html")
            
            # Convert paths to absolute paths
            word_file_abs = os.path.abspath(self.word_file_path)
            pdf_file_abs = os.path.abspath(self.pdf_file_path)
            
            # Get coordinates for highlighting
            x0, y0, x1, y1 = self.pdf_second_page_first_word_coords
            
            # Create HTML content with PDF.js for PDF viewing
            html_content = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <title>Document Link Viewer</title>
                <script src="https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.4.120/pdf.min.js"></script>
                <style>
                    body {{ font-family: Arial, sans-serif; margin: 0; padding: 0; }}
                    .container {{ display: flex; height: 100vh; }}
                    .word-section, .pdf-section {{ flex: 1; padding: 20px; overflow: auto; }}
                    .word-section {{ background-color: #f0f0f0; }}
                    .pdf-section {{ background-color: #e6f2ff; position: relative; }}
                    h2 {{ color: #333; }}
                    .linked-word {{ color: blue; text-decoration: underline; cursor: pointer; }}
                    #pdf-container {{ width: 100%; height: 90%; }}
                    .highlight-overlay {{ 
                        position: absolute; 
                        background-color: rgba(255, 255, 0, 0.5); 
                        border: 2px solid #FF0000;
                        pointer-events: none;
                        display: none;
                    }}
                </style>
            </head>
            <body>
                <div class="container">
                    <div class="word-section">
                        <h2>Word Document</h2>
                        <div id="word-content">
            """
            
            # Add Word content with linked first word
            doc = docx.Document(self.word_file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            for i, para in enumerate(paragraphs):
                if i == 0 and self.word_first_word in para:
                    # Link the first word
                    linked_para = para.replace(
                        self.word_first_word, 
                        f'<span class="linked-word" onclick="goToPdfWord()">{self.word_first_word}</span>',
                        1  # Replace only the first occurrence
                    )
                    html_content += f"<p>{linked_para}</p>"
                else:
                    html_content += f"<p>{para}</p>"
            
            html_content += """
                        </div>
                    </div>
                    <div class="pdf-section">
                        <h2>PDF Document</h2>
                        <div id="pdf-container"></div>
                        <div id="highlight-overlay" class="highlight-overlay"></div>
                    </div>
                </div>
                
                <script>
                    // Initialize PDF.js
                    pdfjsLib.GlobalWorkerOptions.workerSrc = 'https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.4.120/pdf.worker.min.js';
                    
                    let pdfDoc = null;
                    let pdfPageRendering = false;
                    let pageNumPending = null;
                    const scale = 1.5;
                    let pdfCanvas = null;
                    let pdfContext = null;
                    
                    // Load the PDF
                    const loadPdf = async () => {
                        const loadingTask = pdfjsLib.getDocument('file://""" + pdf_file_abs.replace('\\', '\\\\') + """');
                        pdfDoc = await loadingTask.promise;
                        renderPage(1);
                    };
                    
                    // Render the page
                    const renderPage = async (pageNum) => {
                        pdfPageRendering = true;
                        
                        // Get page
                        const page = await pdfDoc.getPage(pageNum);
                        
                        // Create canvas if needed
                        if (!pdfCanvas) {
                            pdfCanvas = document.createElement('canvas');
                            pdfContainer.appendChild(pdfCanvas);
                            pdfContext = pdfCanvas.getContext('2d');
                        }
                        
                        // Calculate viewport
                        const viewport = page.getViewport({ scale });
                        pdfCanvas.height = viewport.height;
                        pdfCanvas.width = viewport.width;
                        
                        // Render PDF page
                        const renderContext = {
                            canvasContext: pdfContext,
                            viewport: viewport
                        };
                        
                        const renderTask = page.render(renderContext);
                        await renderTask.promise;
                        
                        pdfPageRendering = false;
                        
                        // Check if there's a pending page
                        if (pageNumPending !== null) {
                            renderPage(pageNumPending);
                            pageNumPending = null;
                        }
                    };
                    
                    // Go to page
                    const queueRenderPage = (pageNum) => {
                        if (pdfPageRendering) {
                            pageNumPending = pageNum;
                        } else {
                            renderPage(pageNum);
                        }
                    };
                    
                    // Go to the PDF word and highlight it
                    const goToPdfWord = async () => {
                        // Go to page 2
                        queueRenderPage(2);
                        
                        // Add highlight
                        const overlay = document.getElementById('highlight-overlay');
                        
                        // Wait for page to render
                        const checkRendering = () => {
                            if (pdfPageRendering) {
                                setTimeout(checkRendering, 100);
                                return;
                            }
                            
                            // Get the scaling factor
                            const viewport = pdfDoc.getPage(2).getViewport({ scale });
                            const canvas = document.querySelector('canvas');
                            
                            // Calculate position for highlight based on PDF coordinates
                            const x0 = """ + str(x0) + """ * scale;
                            const y0 = """ + str(y0) + """ * scale;
                            const x1 = """ + str(x1) + """ * scale;
                            const y1 = """ + str(y1) + """ * scale;
                            
                            // Get canvas position
                            const rect = canvas.getBoundingClientRect();
                            
                            // Position the highlight
                            overlay.style.left = (x0 + canvas.offsetLeft) + 'px';
                            overlay.style.top = (viewport.height - y1 + canvas.offsetTop) + 'px';
                            overlay.style.width = (x1 - x0) + 'px';
                            overlay.style.height = (y1 - y0) + 'px';
                            overlay.style.display = 'block';
                            
                            // Scroll to the word
                            overlay.scrollIntoView({
                                behavior: 'smooth',
                                block: 'center'
                            });
                        };
                        
                        checkRendering();
                    };
                    
                    // Get reference to the container
                    const pdfContainer = document.getElementById('pdf-container');
                    
                    // Load the PDF when the page loads
                    loadPdf();
                </script>
            </body>
            </html>
            """
            
            # Write to file
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            # Open in browser
            webbrowser.open('file://' + html_path)
            self.status_var.set("Linked document viewer created and opened in browser")
            
        except Exception as e:
            messagebox.showerror("Error", f"Error creating linked viewer: {str(e)}")


if __name__ == "__main__":
    root = tk.Tk()
    app = DocumentLinker(root)
    root.mainloop()