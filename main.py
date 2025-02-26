import os
import sys
import json
from docx import Document
from docx.shared import RGBColor
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import win32com.client
import pythoncom
import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter import ttk
import fitz  # PyMuPDF
import tempfile
import webbrowser
import base64
import re

class WordPdfLinker:
    def __init__(self):
        self.word_path = None
        self.pdf_path = None
        self.mismatches = []
        self.temp_html = None
        self.pdf_document = None
        self.word_document = None
        
    def load_mismatches(self, mismatches_json_path):
        """Load mismatches data from a JSON file."""
        try:
            with open(mismatches_json_path, 'r') as f:
                self.mismatches = json.load(f)
            print(f"Loaded {len(self.mismatches)} mismatches from {mismatches_json_path}")
            return True
        except Exception as e:
            print(f"Error loading mismatches: {e}")
            return False
    
    def highlight_word_document(self, output_path=None):
        """Highlight mismatched words in the Word document and add hyperlinks."""
        try:
            # Load the Word document
            doc = Document(self.word_path)
            
            # Create a unique ID for each mismatch
            for i, mismatch in enumerate(self.mismatches):
                mismatch['id'] = f"mismatch_{i}"
            
            # Process each paragraph in the document
            for para_idx, para in enumerate(doc.paragraphs):
                if not para.text.strip():
                    continue
                
                # Find mismatches in this paragraph
                para_mismatches = [m for m in self.mismatches 
                                  if m['word_location']['paragraph'] == para_idx]
                
                if not para_mismatches:
                    continue
                
                # Sort mismatches by their position (in reverse order to avoid invalidating indices)
                para_mismatches.sort(key=lambda m: m['word_location']['offset'], reverse=True)
                
                # Create a new run for each mismatch
                for mismatch in para_mismatches:
                    word = mismatch['text']
                    offset = mismatch['word_location']['offset']
                    
                    # Find the run containing the word
                    current_offset = 0
                    target_run_idx = None
                    word_start_in_run = None
                    
                    for run_idx, run in enumerate(para.runs):
                        if current_offset <= offset < current_offset + len(run.text):
                            target_run_idx = run_idx
                            word_start_in_run = offset - current_offset
                            break
                        current_offset += len(run.text)
                    
                    if target_run_idx is None:
                        continue
                    
                    # Get the target run
                    run = para.runs[target_run_idx]
                    
                    # Split the run text
                    before_text = run.text[:word_start_in_run]
                    word_text = run.text[word_start_in_run:word_start_in_run + len(word)]
                    after_text = run.text[word_start_in_run + len(word):]
                    
                    # Update the run with the 'before' text
                    run.text = before_text
                    
                    # Create a new run with the word text and apply highlighting
                    new_run = para.add_run(word_text)
                    new_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    
                    # Add bookmark
                    self.add_bookmark(new_run._element, mismatch['id'])
                    
                    # Add hyperlink behavior
                    self.add_hyperlink_style(new_run._element, f"mismatch:{mismatch['id']}")
                    
                    # Add 'after' text in a new run
                    if after_text:
                        para.add_run(after_text)
            
            # Save the document
            save_path = output_path or self.word_path.replace('.docx', '_linked.docx')
            doc.save(save_path)
            print(f"Saved highlighted document to {save_path}")
            return save_path
        
        except Exception as e:
            print(f"Error highlighting Word document: {e}")
            return None
    
    def add_bookmark(self, element, bookmark_id):
        """Add a bookmark to a run element."""
        bookmark_start = OxmlElement('w:bookmarkStart')
        bookmark_start.set(qn('w:id'), '0')
        bookmark_start.set(qn('w:name'), bookmark_id)
        
        bookmark_end = OxmlElement('w:bookmarkEnd')
        bookmark_end.set(qn('w:id'), '0')
        
        element.append(bookmark_start)
        element.append(bookmark_end)
    
    def add_hyperlink_style(self, element, url):
        """Add hyperlink style to a run element."""
        # Add the w:rStyle element
        rStyle = OxmlElement('w:rStyle')
        rStyle.set(qn('w:val'), 'Hyperlink')
        element.append(rStyle)
        
        # Add custom XML properties for the hyperlink
        custom_xml = OxmlElement('w:customXml')
        custom_xml.set(qn('w:uri'), 'link')
        custom_xml.set(qn('w:element'), 'http://schemas.microsoft.com/office/word/2010/wordml')
        
        custom_prop = OxmlElement('w:attr')
        custom_prop.set(qn('w:name'), 'target')
        custom_prop.set(qn('w:val'), url)
        
        custom_xml.append(custom_prop)
        element.append(custom_xml)
    
    def create_pdf_viewer_html(self):
        """Create an HTML file with a PDF viewer that can navigate to specific words."""
        try:
            # Create a temporary HTML file
            fd, html_path = tempfile.mkstemp(suffix='.html')
            os.close(fd)
            
            # Extract the PDF file name
            pdf_filename = os.path.basename(self.pdf_path)
            
            # Create the HTML content with PDF.js
            html_content = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <title>PDF Viewer</title>
                <script src="https://cdnjs.cloudflare.com/ajax/libs/pdf.js/2.12.313/pdf.min.js"></script>
                <style>
                    body {{
                        margin: 0;
                        padding: 0;
                        display: flex;
                        flex-direction: column;
                        height: 100vh;
                    }}
                    #toolbar {{
                        padding: 10px;
                        background-color: #f0f0f0;
                        border-bottom: 1px solid #ddd;
                    }}
                    #viewer {{
                        flex-grow: 1;
                        overflow: auto;
                    }}
                    #status {{
                        padding: 5px 10px;
                        background-color: #e0e0e0;
                        font-size: 14px;
                    }}
                    .highlight {{
                        background-color: yellow;
                        border-radius: 3px;
                        padding: 2px;
                    }}
                </style>
            </head>
            <body>
                <div id="toolbar">
                    <button id="prev">Previous</button>
                    <button id="next">Next</button>
                    <span>Page: <span id="page_num"></span> / <span id="page_count"></span></span>
                    <span id="word-info" style="margin-left: 20px;"></span>
                </div>
                <div id="viewer"></div>
                <div id="status">Ready</div>
                
                <script>
                    // Load PDF.js
                    pdfjsLib.GlobalWorkerOptions.workerSrc = 'https://cdnjs.cloudflare.com/ajax/libs/pdf.js/2.12.313/pdf.worker.min.js';
                    
                    // The mismatches data
                    const mismatches = {JSON_MISMATCHES};
                    
                    // Current selected mismatch
                    let selectedMismatchId = null;
                    
                    // Initialize the PDF viewer
                    let pdfDoc = null;
                    let pageNum = 1;
                    let pageRendering = false;
                    let pageNumPending = null;
                    let scale = 1.5;
                    let canvas = document.createElement('canvas');
                    let ctx = canvas.getContext('2d');
                    let viewer = document.getElementById('viewer');
                    viewer.appendChild(canvas);
                    
                    // Load the PDF
                    const loadPdf = async () => {{
                        try {{
                            const loadingTask = pdfjsLib.getDocument({{
                                url: '{pdf_filename}',
                                cMapUrl: 'https://cdnjs.cloudflare.com/ajax/libs/pdf.js/2.12.313/cmaps/',
                                cMapPacked: true,
                            }});
                            
                            pdfDoc = await loadingTask.promise;
                            document.getElementById('page_count').textContent = pdfDoc.numPages;
                            
                            // Initial render
                            renderPage(pageNum);
                        }} catch (error) {{
                            console.error('Error loading PDF:', error);
                            document.getElementById('status').textContent = 'Error loading PDF: ' + error.message;
                        }}
                    }};
                    
                    // Render a page
                    const renderPage = async (num) => {{
                        pageRendering = true;
                        document.getElementById('page_num').textContent = num;
                        
                        try {{
                            const page = await pdfDoc.getPage(num);
                            const viewport = page.getViewport({{ scale }});
                            
                            canvas.height = viewport.height;
                            canvas.width = viewport.width;
                            
                            const renderContext = {{
                                canvasContext: ctx,
                                viewport: viewport
                            }};
                            
                            await page.render(renderContext).promise;
                            
                            // Extract text content for highlighting
                            const textContent = await page.getTextContent();
                            
                            // Find mismatches on this page
                            const pageMismatches = mismatches.filter(m => m.pdf_location.page === num);
                            
                            if (pageMismatches.length > 0) {{
                                // Create text overlay for highlighting
                                const textLayerDiv = document.createElement('div');
                                textLayerDiv.className = 'textLayer';
                                textLayerDiv.style.position = 'absolute';
                                textLayerDiv.style.left = canvas.offsetLeft + 'px';
                                textLayerDiv.style.top = canvas.offsetTop + 'px';
                                textLayerDiv.style.height = canvas.height + 'px';
                                textLayerDiv.style.width = canvas.width + 'px';
                                viewer.appendChild(textLayerDiv);
                                
                                // Create text spans
                                let lastSpan = null;
                                let lastY = null;
                                
                                for (const item of textContent.items) {{
                                    const tx = pdfjsLib.Util.transform(
                                        viewport.transform,
                                        [1, 0, 0, -1, item.transform[4], item.transform[5]]
                                    );
                                    
                                    const style = window.getComputedStyle(textLayerDiv);
                                    const fontAscent = parseFloat(style.fontSize) * 0.8;
                                    
                                    const div = document.createElement('span');
                                    div.textContent = item.str;
                                    div.style.left = Math.floor(tx[0]) + 'px';
                                    div.style.top = Math.floor(tx[1] - fontAscent) + 'px';
                                    div.style.fontSize = Math.floor(item.height * viewport.scale) + 'px';
                                    div.style.fontFamily = item.fontName;
                                    div.style.position = 'absolute';
                                    
                                    // Check if this text item matches any of our mismatches
                                    for (const mismatch of pageMismatches) {{
                                        if (item.str.includes(mismatch.text)) {{
                                            const matchIndex = item.str.indexOf(mismatch.text);
                                            
                                            if (matchIndex !== -1) {{
                                                // Create a wrapper span for the word
                                                const wordSpan = document.createElement('span');
                                                wordSpan.textContent = mismatch.text;
                                                wordSpan.className = 'highlight';
                                                wordSpan.dataset.mismatchId = mismatch.id;
                                                
                                                // Mark as selected if this is the current mismatch
                                                if (mismatch.id === selectedMismatchId) {{
                                                    wordSpan.style.backgroundColor = 'green';
                                                    wordSpan.style.color = 'white';
                                                    
                                                    // Scroll this item into view
                                                    setTimeout(() => {{
                                                        wordSpan.scrollIntoView({{
                                                            behavior: 'smooth',
                                                            block: 'center'
                                                        }});
                                                        
                                                        // Update info
                                                        const wordInfo = document.getElementById('word-info');
                                                        wordInfo.textContent = `Viewing: "${mismatch.text}" - PDF formatting: ${formatFormattingInfo(mismatch.pdf_formatting)}, Word formatting: ${formatFormattingInfo(mismatch.word_formatting)}`;
                                                    }}, 100);
                                                }}
                                                
                                                // Handle click to select this mismatch
                                                wordSpan.addEventListener('click', () => {{
                                                    selectedMismatchId = mismatch.id;
                                                    renderPage(pageNum);
                                                }});
                                                
                                                div.textContent = item.str.substring(0, matchIndex);
                                                div.appendChild(wordSpan);
                                                div.appendChild(document.createTextNode(item.str.substring(matchIndex + mismatch.text.length)));
                                                break;
                                            }}
                                        }}
                                    }}
                                    
                                    textLayerDiv.appendChild(div);
                                }}
                            }}
                            
                            pageRendering = false;
                            if (pageNumPending !== null) {{
                                renderPage(pageNumPending);
                                pageNumPending = null;
                            }}
                        }} catch (error) {{
                            console.error('Error rendering page:', error);
                            document.getElementById('status').textContent = 'Error rendering page: ' + error.message;
                            pageRendering = false;
                        }}
                    }};
                    
                    // Format the formatting info for display
                    function formatFormattingInfo(formatting) {{
                        const styles = [];
                        if (formatting.bold) styles.push('Bold');
                        if (formatting.italic) styles.push('Italic');
                        if (formatting.underline) styles.push('Underline');
                        return styles.join(', ') || 'Normal';
                    }}
                    
                    // Go to previous page
                    document.getElementById('prev').addEventListener('click', () => {{
                        if (pageNum <= 1) return;
                        pageNum--;
                        queueRenderPage(pageNum);
                    }});
                    
                    // Go to next page
                    document.getElementById('next').addEventListener('click', () => {{
                        if (pageNum >= pdfDoc.numPages) return;
                        pageNum++;
                        queueRenderPage(pageNum);
                    }});
                    
                    // Queue rendering of a page
                    function queueRenderPage(num) {{
                        if (pageRendering) {{
                            pageNumPending = num;
                        }} else {{
                            renderPage(num);
                        }}
                    }}
                    
                    // Handle navigation from Word
                    function navigateToMismatch(mismatchId) {{
                        const mismatch = mismatches.find(m => m.id === mismatchId);
                        if (mismatch) {{
                            selectedMismatchId = mismatchId;
                            pageNum = mismatch.pdf_location.page;
                            renderPage(pageNum);
                        }}
                    }}
                    
                    // Check for mismatch ID in URL
                    window.onload = () => {{
                        loadPdf();
                        
                        const urlParams = new URLSearchParams(window.location.search);
                        const mismatchId = urlParams.get('mismatch');
                        if (mismatchId) {{
                            navigateToMismatch(mismatchId);
                        }}
                    }};
                </script>
            </body>
            </html>
            """.replace("{JSON_MISMATCHES}", json.dumps(self.mismatches))
            
            # Write the HTML to the temporary file
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            # Copy the PDF to the same directory as the HTML
            pdf_copy_path = os.path.join(os.path.dirname(html_path), pdf_filename)
            with open(self.pdf_path, 'rb') as src, open(pdf_copy_path, 'wb') as dst:
                dst.write(src.read())
            
            self.temp_html = html_path
            return html_path
        
        except Exception as e:
            print(f"Error creating PDF viewer HTML: {e}")
            return None
    
    def setup_word_event_handler(self, word_doc_path):
        """Set up an event handler for the Word document to handle clicks on mismatched words."""
        try:
            # Initialize COM
            pythoncom.CoInitialize()
            
            # Create a Word application
            word_app = win32com.client.Dispatch("Word.Application")
            word_app.Visible = True
            
            # Open the document
            word_doc = word_app.Documents.Open(word_doc_path)
            
            # Create a temporary VBA module to handle the hyperlink event
            vba_code = f"""
            Sub HyperlinkHandler()
                Dim hyperlink As String
                hyperlink = Selection.Hyperlinks(1).Address
                
                If Left(hyperlink, 9) = "mismatch:" Then
                    Dim mismatchId As String
                    mismatchId = Mid(hyperlink, 10)
                    
                    ' Open the PDF viewer with the mismatch ID
                    Shell "cmd /c start {self.temp_html}?mismatch=" & mismatchId, vbNormalFocus
                End If
            End Sub
            """
            
            # Add the VBA module
            vba_module = word_doc.VBProject.VBComponents.Add(1)  # 1 = vbext_ct_StdModule
            vba_module.CodeModule.AddFromString(vba_code)
            
            # Set up the document to use the hyperlink handler
            word_app.OnTime(Now(), "HyperlinkHandler")
            
            print("Word event handler set up successfully")
            
            # Keep references to prevent garbage collection
            self.word_app = word_app
            self.word_doc = word_doc
            
            return True
        
        except Exception as e:
            print(f"Error setting up Word event handler: {e}")
            return False
    
    def run_gui(self):
        """Run a GUI to load documents and start the linking process."""
        root = tk.Tk()
        root.title("Word-PDF Linker")
        root.geometry("600x500")
        
        # Create a main frame
        main_frame = ttk.Frame(root, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Create widgets
        ttk.Label(main_frame, text="Word-PDF Format Mismatch Linker", font=("Arial", 16)).pack(pady=10)
        
        # Word document selection
        word_frame = ttk.LabelFrame(main_frame, text="Word Document")
        word_frame.pack(fill=tk.X, pady=5)
        
        word_path_var = tk.StringVar()
        ttk.Entry(word_frame, textvariable=word_path_var, width=50).pack(side=tk.LEFT, padx=5, pady=5, fill=tk.X, expand=True)
        
        def browse_word():
            file_path = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
            if file_path:
                word_path_var.set(file_path)
                self.word_path = file_path
        
        ttk.Button(word_frame, text="Browse", command=browse_word).pack(side=tk.RIGHT, padx=5, pady=5)
        
        # PDF document selection
        pdf_frame = ttk.LabelFrame(main_frame, text="PDF Document")
        pdf_frame.pack(fill=tk.X, pady=5)
        
        pdf_path_var = tk.StringVar()
        ttk.Entry(pdf_frame, textvariable=pdf_path_var, width=50).pack(side=tk.LEFT, padx=5, pady=5, fill=tk.X, expand=True)
        
        def browse_pdf():
            file_path = filedialog.askopenfilename(filetypes=[("PDF Documents", "*.pdf")])
            if file_path:
                pdf_path_var.set(file_path)
                self.pdf_path = file_path
        
        ttk.Button(pdf_frame, text="Browse", command=browse_pdf).pack(side=tk.RIGHT, padx=5, pady=5)
        
        # Mismatches JSON selection
        mismatches_frame = ttk.LabelFrame(main_frame, text="Mismatches JSON")
        mismatches_frame.pack(fill=tk.X, pady=5)
        
        mismatches_path_var = tk.StringVar()
        ttk.Entry(mismatches_frame, textvariable=mismatches_path_var, width=50).pack(side=tk.LEFT, padx=5, pady=5, fill=tk.X, expand=True)
        
        def browse_mismatches():
            file_path = filedialog.askopenfilename(filetypes=[("JSON Files", "*.json")])
            if file_path:
                mismatches_path_var.set(file_path)
                self.load_mismatches(file_path)
        
        ttk.Button(mismatches_frame, text="Browse", command=browse_mismatches).pack(side=tk.RIGHT, padx=5, pady=5)
        
        # Status display
        status_frame = ttk.LabelFrame(main_frame, text="Status")
        status_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        
        status_text = tk.Text(status_frame, height=10, width=50)
        status_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        scrollbar = ttk.Scrollbar(status_text, command=status_text.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        status_text.config(yscrollcommand=scrollbar.set)
        
        def log(message):
            status_text.insert(tk.END, message + "\n")
            status_text.see(tk.END)
            print(message)
        
        # Redirect print to the status text
        class TextRedirector:
            def __init__(self, text_widget):
                self.text_widget = text_widget
            
            def write(self, string):
                self.text_widget.insert(tk.END, string)
                self.text_widget.see(tk.END)
            
            def flush(self):
                pass
        
        sys.stdout = TextRedirector(status_text)
        
        # Process button
        def process_documents():
            if not self.word_path:
                messagebox.showerror("Error", "Please select a Word document")
                return
            
            if not self.pdf_path:
                messagebox.showerror("Error", "Please select a PDF document")
                return
            
            if not self.mismatches:
                messagebox.showerror("Error", "Please load mismatches JSON data")
                return
            
            try:
                log("Processing documents...")
                
                # Highlight and add links to Word document
                linked_doc_path = self.highlight_word_document()
                
                if not linked_doc_path:
                    messagebox.showerror("Error", "Failed to process Word document")
                    return
                
                log(f"Created linked Word document: {linked_doc_path}")
                
                # Create PDF viewer HTML
                html_path = self.create_pdf_viewer_html()
                
                if not html_path:
                    messagebox.showerror("Error", "Failed to create PDF viewer")
                    return
                
                log(f"Created PDF viewer HTML: {html_path}")
                
                # Set up Word event handler
                if self.setup_word_event_handler(linked_doc_path):
                    log("Word event handler set up successfully")
                else:
                    log("Failed to set up Word event handler, but you can still use the documents")
                
                log("Processing complete. Open the linked Word document and click on highlighted words to view them in the PDF.")
            
            except Exception as e:
                log(f"Error processing documents: {e}")
                messagebox.showerror("Error", f"Processing failed: {e}")
        
        ttk.Button(main_frame, text="Process Documents", command=process_documents).pack(pady=10)
        
        # Run the GUI
        root.mainloop()

# Example usage
if __name__ == "__main__":
    linker = WordPdfLinker()
    linker.run_gui()