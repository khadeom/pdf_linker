import docx
from docx import Document
import PyPDF2
import re
import os
from docx.shared import Pt
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

def extract_pdf_text(pdf_path):
    """Extract text from a PDF file."""
    with open(pdf_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
    return text

def extract_word_text(docx_path):
    """Extract text from a Word document."""
    doc = Document(docx_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

def find_word_positions(word, text):
    """Find all positions of a word in text."""
    positions = []
    pattern = r'\b' + re.escape(word) + r'\b'
    for match in re.finditer(pattern, text):
        positions.append(match.start())
    return positions

def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Add styling (optional)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')  # Blue color
    rPr.append(color)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')  # Single underline
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    
    return hyperlink

def create_word_links_to_pdf(word_path, pdf_path, output_word_path, target_words):
    """Create links in Word document to PDF for specific words."""
    # Extract text from PDF to find word positions
    pdf_text = extract_pdf_text(pdf_path)
    
    # Create a new Word document based on the original
    doc = Document(word_path)
    
    # PDF file path (for creating links)
    pdf_abs_path = os.path.abspath(pdf_path)
    
    # Process each paragraph in Word
    for para in doc.paragraphs:
        for word in target_words:
            if word in para.text:
                # Find positions in PDF
                pdf_positions = find_word_positions(word, pdf_text)
                if pdf_positions:
                    # Create link to first occurrence in PDF
                    # Format: file:///path/to/file.pdf#page=X&search=word
                    # Note: You might need to calculate the actual page number
                    link_url = f"file:///{pdf_abs_path}#search={word}"
                    
                    # Create a new paragraph with the link
                    new_para = doc.add_paragraph()
                    add_hyperlink(new_para, f"Link to '{word}' in PDF", link_url)
    
    # Save the modified document
    doc.save(output_word_path)

# Example usage
word_path = "document.docx"
pdf_path = "document.pdf"
output_word_path = "linked_document.docx"
target_words = ["important", "keyword", "reference"]

create_word_links_to_pdf(word_path, pdf_path, output_word_path, target_words)