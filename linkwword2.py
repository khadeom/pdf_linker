
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import docx
import PyPDF2
import os
import webbrowser
import tempfile
from pathlib import Path
import re

class DocumentLinker:
    def __init__(self, root):
        self.root = root
        self.root.title("Document Word Linker")
        self.root.geometry("800x600")
        
        # File paths
        self.word_file_path = None
        self.pdf_file_path = None
        
        # Extracted content
        self.word_first_word = None
        self.pdf_second_page_text = None
        self.pdf_second_page_first_word = None
        
        # Setup UI
        self.create_widgets()
    
    def create_widgets(self):
        # Frame for file selection
        file_frame = ttk.LabelFrame(self.root, text="Document Selection")
        file_frame.pack(fill="x", padx=10, pady=10)
        
        # Word document selection
        ttk.Label(file_frame, text="Word Document:").grid(row=0, column=0, padx=5, pady=5, sticky="w")
        self.word_file_label = ttk.Label(file_frame, text="No file selected")
        self.word_file_label.grid(row=0, column=1, padx=5, pady=5, sticky="w")
        ttk.Button(file_frame, text="Browse...", command=self.select_word_file).grid(row=0, column=2, padx=5, pady=5)
        
        # PDF document selection
        ttk.Label(file_frame, text="PDF Document:").grid(row=1, column=0, padx=5, pady=5, sticky="w")
        self.pdf_file_label = ttk.Label(file_frame, text="No file selected")
        self.pdf_file_label.grid(row=1, column=1, padx=5, pady=5, sticky="w")
        ttk.Button(file_frame, text="Browse...", command=self.select_pdf_file).grid(row=1, column=2, padx=5, pady=5)
        
        # Frame for process and analysis
        analysis_frame = ttk.LabelFrame(self.root, text="Document Analysis")
        analysis_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        # Word document first word display
        ttk.Label(analysis_frame, text="First word in Word document:").grid(row=0, column=0, padx=5, pady=5, sticky="w")
        self.word_first_word_label = ttk.Label(analysis_frame, text="-")
        self.word_first_word_label.grid(row=0, column=1, padx=5, pady=5, sticky="w")
        
        # PDF second page first word display
        ttk.Label(analysis_frame, text="First word on second page of PDF:").grid(row=1, column=0, padx=5, pady=5, sticky="w")
        self.pdf_first_word_label = ttk.Label(analysis_frame, text="-")
        self.pdf_first_word_label.grid(row=1, column=1, padx=5, pady=5, sticky="w")
        
        # Preview frame
        preview_frame = ttk.LabelFrame(self.root, text="Preview")
        preview_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        # Text preview for Word
        ttk.Label(preview_frame, text="Word Document Preview:").grid(row=0, column=0, padx=5, pady=5, sticky="nw")
        self.word_preview = tk.Text(preview_frame, height=5, width=80, wrap=tk.WORD)
        self.word_preview.grid(row=1, column=0, padx=5, pady=5, sticky="nsew")
        word_scroll = ttk.Scrollbar(preview_frame, orient="vertical", command=self.word_preview.yview)
        word_scroll.grid(row=1, column=1, sticky="ns")
        self.word_preview.config(yscrollcommand=word_scroll.set)
        
        # Text preview for PDF
        ttk.Label(preview_frame, text="PDF Second Page Preview:").grid(row=2, column=0, padx=5, pady=5, sticky="nw")
        self.pdf_preview = tk.Text(preview_frame, height=5, width=80, wrap=tk.WORD)
        self.pdf_preview.grid(row=3, column=0, padx=5, pady=5, sticky="nsew")
        pdf_scroll = ttk.Scrollbar(preview_frame, orient="vertical", command=self.pdf_preview.yview)
        pdf_scroll.grid(row=3, column=1, sticky="ns")
        self.pdf_preview.config(yscrollcommand=pdf_scroll.set)
        
        # Configure row and column weights for proper resizing
        preview_frame.columnconfigure(0, weight=1)
        preview_frame.rowconfigure(1, weight=1)
        preview_frame.rowconfigure(3, weight=1)
        
        # Buttons
        button_frame = ttk.Frame(self.root)
        button_frame.pack(fill="x", padx=10, pady=10)
        
        ttk.Button(button_frame, text="Process Documents", command=self.process_documents).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Create Linked HTML", command=self.create_linked_html).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Exit", command=self.root.quit).pack(side=tk.RIGHT, padx=5)
        
        # Status bar
        self.status_var = tk.StringVar()
        self.status_var.set("Ready")
        status_bar = ttk.Label(self.root, textvariable=self.status_var, relief=tk.SUNKEN, anchor=tk.W)
        status_bar.pack(side=tk.BOTTOM, fill=tk.X)
    
    def select_word_file(self):
        file_path = filedialog.askopenfilename(
            title="Select Word Document",
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
        )
        if file_path:
            self.word_file_path = file_path
            self.word_file_label.config(text=os.path.basename(file_path))
            self.status_var.set(f"Word document selected: {os.path.basename(file_path)}")
    
    def select_pdf_file(self):
        file_path = filedialog.askopenfilename(
            title="Select PDF Document",
            filetypes=[("PDF Documents", "*.pdf"), ("All Files", "*.*")]
        )
        if file_path:
            self.pdf_file_path = file_path
            self.pdf_file_label.config(text=os.path.basename(file_path))
            self.status_var.set(f"PDF document selected: {os.path.basename(file_path)}")
    
    def extract_word_content(self):
        try:
            doc = docx.Document(self.word_file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            if not paragraphs:
                messagebox.showwarning("Warning", "No text found in Word document")
                return False
            
            # Get the first word from first paragraph
            first_paragraph = paragraphs[0]
            first_word = first_paragraph.split()[0] if first_paragraph.split() else ""
            
            self.word_first_word = first_word
            self.word_first_word_label.config(text=first_word)
            
            # Display preview
            self.word_preview.delete(1.0, tk.END)
            preview_text = "\n\n".join(paragraphs[:3])  # First 3 paragraphs
            self.word_preview.insert(tk.END, preview_text)
            
            return True
        except Exception as e:
            messagebox.showerror("Error", f"Error extracting Word content: {str(e)}")
            return False
    
    def extract_pdf_content(self):
        try:
            reader = PyPDF2.PdfReader(self.pdf_file_path)
            
            if len(reader.pages) < 2:
                messagebox.showwarning("Warning", "PDF has fewer than 2 pages")
                return False
            
            # Get text from second page
            page = reader.pages[1]  # 0-based index, so 1 is second page
            page_text = page.extract_text()
            
            if not page_text.strip():
                messagebox.showwarning("Warning", "No text found on second page of PDF")
                return False
            
            # Get first word from second page
            first_word = page_text.split()[0] if page_text.split() else ""
            
            self.pdf_second_page_text = page_text
            self.pdf_second_page_first_word = first_word
            self.pdf_first_word_label.config(text=first_word)
            
            # Display preview
            self.pdf_preview.delete(1.0, tk.END)
            self.pdf_preview.insert(tk.END, page_text[:500] + "..." if len(page_text) > 500 else page_text)
            
            return True
        except Exception as e:
            messagebox.showerror("Error", f"Error extracting PDF content: {str(e)}")
            return False
    
    def process_documents(self):
        if not self.word_file_path or not self.pdf_file_path:
            messagebox.showwarning("Warning", "Please select both Word and PDF documents")
            return
        
        self.status_var.set("Processing documents...")
        self.root.update_idletasks()
        
        word_success = self.extract_word_content()
        pdf_success = self.extract_pdf_content()
        
        if word_success and pdf_success:
            self.status_var.set(f"Linking '{self.word_first_word}' to '{self.pdf_second_page_first_word}'")
            messagebox.showinfo("Success", "Documents processed successfully. Ready to create linked view.")
        else:
            self.status_var.set("Processing failed")
    
    def create_linked_html(self):
        if not self.word_first_word or not self.pdf_second_page_text:
            messagebox.showwarning("Warning", "Please process documents first")
            return
        
        try:
            # Create temporary HTML file
            temp_dir = tempfile.gettempdir()
            html_path = os.path.join(temp_dir, "document_link.html")
            
            # Create HTML content with linking
            html_content = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <title>Document Link Viewer</title>
                <style>
                    body {{ font-family: Arial, sans-serif; margin: 0; padding: 0; }}
                    .container {{ display: flex; height: 100vh; }}
                    .word-section, .pdf-section {{ flex: 1; padding: 20px; overflow: auto; }}
                    .word-section {{ background-color: #f0f0f0; }}
                    .pdf-section {{ background-color: #e6f2ff; }}
                    h2 {{ color: #333; }}
                    .highlight {{ background-color: yellow; font-weight: bold; }}
                    .linked-word {{ color: blue; text-decoration: underline; cursor: pointer; }}
                </style>
                <script>
                    function scrollToPdfWord() {{
                        document.getElementById('pdf-target').scrollIntoView({{
                            behavior: 'smooth'
                        }});
                    }}
                </script>
            </head>
            <body>
                <div class="container">
                    <div class="word-section">
                        <h2>Word Document</h2>
                        <p>
            """
            
            # Add Word content with linked first word
            doc = docx.Document(self.word_file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            for i, para in enumerate(paragraphs):
                if i == 0 and self.word_first_word in para:
                    # Link the first word
                    linked_para = para.replace(
                        self.word_first_word, 
                        f'<span class="linked-word" onclick="scrollToPdfWord()">{self.word_first_word}</span>',
                        1  # Replace only the first occurrence
                    )
                    html_content += linked_para + "</p><p>"
                else:
                    html_content += para + "</p><p>"
            
            html_content += """
                        </p>
                    </div>
                    <div class="pdf-section">
                        <h2>PDF Document (Page 2)</h2>
                        <p>
            """
            
            # Add PDF content with highlighted first word
            pdf_text = self.pdf_second_page_text
            if self.pdf_second_page_first_word in pdf_text:
                # Highlight the first word
                highlighted_text = pdf_text.replace(
                    self.pdf_second_page_first_word,
                    f'<span id="pdf-target" class="highlight">{self.pdf_second_page_first_word}</span>',
                    1  # Replace only the first occurrence
                )
                html_content += highlighted_text.replace("\n", "<br>")
            else:
                html_content += pdf_text.replace("\n", "<br>")
            
            html_content += """
                        </p>
                    </div>
                </div>
            </body>
            </html>
            """
            
            # Write to file
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            # Open in browser
            webbrowser.open('file://' + html_path)
            self.status_var.set("Linked HTML document created and opened in browser")
            
        except Exception as e:
            messagebox.showerror("Error", f"Error creating linked view: {str(e)}")


if __name__ == "__main__":
    root = tk.Tk()
    app = DocumentLinker(root)
    root.mainloop()