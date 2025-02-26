
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import docx
import os
import webbrowser
import tempfile
import shutil
import fitz  # PyMuPDF
import base64

class DocumentLinker:
    def __init__(self, root):
        self.root = root
        self.root.title("Document Word Linker")
        self.root.geometry("800x600")
        
        # File paths
        self.word_file_path = None
        self.pdf_file_path = None
        
        # Extracted content
        self.word_first_word = None
        self.pdf_second_page_first_word = None
        self.pdf_second_page_first_word_coords = None
        
        # Setup UI
        self.create_widgets()
    
    def create_widgets(self):
        # Frame for file selection
        file_frame = ttk.LabelFrame(self.root, text="Document Selection")
        file_frame.pack(fill="x", padx=10, pady=10)
        
        # Word document selection
        ttk.Label(file_frame, text="Word Document:").grid(row=0, column=0, padx=5, pady=5, sticky="w")
        self.word_file_label = ttk.Label(file_frame, text="No file selected")
        self.word_file_label.grid(row=0, column=1, padx=5, pady=5, sticky="w")
        ttk.Button(file_frame, text="Browse...", command=self.select_word_file).grid(row=0, column=2, padx=5, pady=5)
        
        # PDF document selection
        ttk.Label(file_frame, text="PDF Document:").grid(row=1, column=0, padx=5, pady=5, sticky="w")
        self.pdf_file_label = ttk.Label(file_frame, text="No file selected")
        self.pdf_file_label.grid(row=1, column=1, padx=5, pady=5, sticky="w")
        ttk.Button(file_frame, text="Browse...", command=self.select_pdf_file).grid(row=1, column=2, padx=5, pady=5)
        
        # Frame for process and analysis
        analysis_frame = ttk.LabelFrame(self.root, text="Document Analysis")
        analysis_frame.pack(fill="x", padx=10, pady=10)
        
        # Word document first word display
        ttk.Label(analysis_frame, text="First word in Word document:").grid(row=0, column=0, padx=5, pady=5, sticky="w")
        self.word_first_word_label = ttk.Label(analysis_frame, text="-")
        self.word_first_word_label.grid(row=0, column=1, padx=5, pady=5, sticky="w")
        
        # PDF second page first word display
        ttk.Label(analysis_frame, text="First word on second page of PDF:").grid(row=1, column=0, padx=5, pady=5, sticky="w")
        self.pdf_first_word_label = ttk.Label(analysis_frame, text="-")
        self.pdf_first_word_label.grid(row=1, column=1, padx=5, pady=5, sticky="w")
        
        # Buttons
        button_frame = ttk.Frame(self.root)
        button_frame.pack(fill="x", padx=10, pady=10)
        
        ttk.Button(button_frame, text="Process Documents", command=self.process_documents).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Create Linked Viewer", command=self.create_linked_viewer).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Exit", command=self.root.quit).pack(side=tk.RIGHT, padx=5)
        
        # Status bar
        self.status_var = tk.StringVar()
        self.status_var.set("Ready")
        status_bar = ttk.Label(self.root, textvariable=self.status_var, relief=tk.SUNKEN, anchor=tk.W)
        status_bar.pack(side=tk.BOTTOM, fill=tk.X)
    
    def select_word_file(self):
        file_path = filedialog.askopenfilename(
            title="Select Word Document",
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
        )
        if file_path:
            self.word_file_path = file_path
            self.word_file_label.config(text=os.path.basename(file_path))
            self.status_var.set(f"Word document selected: {os.path.basename(file_path)}")
    
    def select_pdf_file(self):
        file_path = filedialog.askopenfilename(
            title="Select PDF Document",
            filetypes=[("PDF Documents", "*.pdf"), ("All Files", "*.*")]
        )
        if file_path:
            self.pdf_file_path = file_path
            self.pdf_file_label.config(text=os.path.basename(file_path))
            self.status_var.set(f"PDF document selected: {os.path.basename(file_path)}")
    
    def extract_word_content(self):
        try:
            doc = docx.Document(self.word_file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            if not paragraphs:
                messagebox.showwarning("Warning", "No text found in Word document")
                return False
            
            # Get the first word from first paragraph
            first_paragraph = paragraphs[0]
            first_word = first_paragraph.split()[0] if first_paragraph.split() else ""
            
            self.word_first_word = first_word
            self.word_first_word_label.config(text=first_word)
            
            return True
        except Exception as e:
            messagebox.showerror("Error", f"Error extracting Word content: {str(e)}")
            return False
    
    def extract_pdf_content(self):
        try:
            # Using PyMuPDF (fitz) to extract text and positions
            pdf_document = fitz.open(self.pdf_file_path)
            
            if pdf_document.page_count < 2:
                messagebox.showwarning("Warning", "PDF has fewer than 2 pages")
                return False
            
            # Get second page
            page = pdf_document[1]  # 0-based index, so 1 is the second page
            
            # Get first word from second page and its coordinates
            words = page.get_text("words")  # List of (x0, y0, x1, y1, word, block_no, line_no, word_no)
            
            if not words:
                messagebox.showwarning("Warning", "No words found on second page of PDF")
                return False
            
            first_word_info = words[0]
            first_word = first_word_info[4]  # The word is at index 4
            word_coords = first_word_info[:4]  # Coordinates are at indices 0-3 (x0,y0,x1,y1)
            
            self.pdf_second_page_first_word = first_word
            self.pdf_second_page_first_word_coords = word_coords
            self.pdf_first_word_label.config(text=first_word)
            
            return True
        except Exception as e:
            messagebox.showerror("Error", f"Error extracting PDF content: {str(e)}")
            return False
    
    def process_documents(self):
        if not self.word_file_path or not self.pdf_file_path:
            messagebox.showwarning("Warning", "Please select both Word and PDF documents")
            return
        
        self.status_var.set("Processing documents...")
        self.root.update_idletasks()
        
        word_success = self.extract_word_content()
        pdf_success = self.extract_pdf_content()
        
        if word_success and pdf_success:
            self.status_var.set(f"Linking '{self.word_first_word}' to '{self.pdf_second_page_first_word}'")
            messagebox.showinfo("Success", "Documents processed successfully. Ready to create linked viewer.")
        else:
            self.status_var.set("Processing failed")
    
    def create_linked_viewer(self):
        if not self.word_first_word or not self.pdf_second_page_first_word:
            messagebox.showwarning("Warning", "Please process documents first")
            return
        
        try:
            # Create temp directory for viewer files
            temp_dir = tempfile.mkdtemp()
            html_path = os.path.join(temp_dir, "viewer.html")
            
            # Copy PDF to temp directory
            pdf_temp_path = os.path.join(temp_dir, "document.pdf")
            shutil.copy2(self.pdf_file_path, pdf_temp_path)
            
            # Extract coordinates
            x0, y0, x1, y1 = self.pdf_second_page_first_word_coords
            
            # Create HTML content with PDF.js
            html_content = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <title>Document Link Viewer</title>
                <script src="https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.4.120/pdf.min.js"></script>
                <style>
                    body {{ font-family: Arial, sans-serif; margin: 0; padding: 0; }}
                    .container {{ display: flex; height: 100vh; }}
                    .word-section, .pdf-section {{ flex: 1; padding: 20px; overflow: auto; box-sizing: border-box; }}
                    .word-section {{ background-color: #f0f0f0; }}
                    .pdf-section {{ background-color: #e6f2ff; position: relative; }}
                    h2 {{ color: #333; }}
                    .linked-word {{ color: blue; text-decoration: underline; cursor: pointer; }}
                    #pdf-viewer {{ width: 100%; height: 90%; border: 1px solid #ccc; }}
                </style>
            </head>
            <body>
                <div class="container">
                    <div class="word-section">
                        <h2>Word Document</h2>
                        <div id="word-content">
            """
            
            # Add Word content with linked first word
            doc = docx.Document(self.word_file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            for i, para in enumerate(paragraphs):
                if i == 0 and self.word_first_word in para:
                    # Link the first word
                    linked_para = para.replace(
                        self.word_first_word, 
                        f'<span class="linked-word" onclick="goToPdfWord()">{self.word_first_word}</span>',
                        1  # Replace only the first occurrence
                    )
                    html_content += f"<p>{linked_para}</p>"
                else:
                    html_content += f"<p>{para}</p>"
            
            html_content += """
                        </div>
                    </div>
                    <div class="pdf-section">
                        <h2>PDF Document</h2>
                        <div id="pdf-viewer"></div>
                    </div>
                </div>
                
                <script>
                    // Initialize PDF.js
                    pdfjsLib.GlobalWorkerOptions.workerSrc = 'https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.4.120/pdf.worker.min.js';
                    
                    // PDF viewer element
                    const pdfViewer = document.getElementById('pdf-viewer');
                    
                    // Load and display the PDF
                    let pdfDoc = null;
                    let pdfPage = null;
                    let pdfCanvas = null;
                    let pdfRenderTask = null;
                    let currentPageNum = 1;
                    const scale = 1.5;
                    
                    // Initialize viewer
                    async function initPdfViewer() {
                        try {
                            // Load the PDF
                            const loadingTask = pdfjsLib.getDocument('document.pdf');
                            pdfDoc = await loadingTask.promise;
                            renderPage(1);
                        } catch (error) {
                            console.error('Error loading PDF:', error);
                            pdfViewer.textContent = 'Error loading PDF: ' + error.message;
                        }
                    }
                    
                    // Render page
                    async function renderPage(pageNum) {
                        if (pdfRenderTask) {
                            // Cancel any pending render task
                            pdfRenderTask.cancel();
                        }
                        
                        try {
                            // Get the page
                            pdfPage = await pdfDoc.getPage(pageNum);
                            currentPageNum = pageNum;
                            
                            // Create canvas if needed
                            if (!pdfCanvas) {
                                pdfCanvas = document.createElement('canvas');
                                pdfViewer.appendChild(pdfCanvas);
                            }
                            
                            // Prepare canvas for rendering
                            const viewport = pdfPage.getViewport({ scale });
                            const context = pdfCanvas.getContext('2d');
                            pdfCanvas.width = viewport.width;
                            pdfCanvas.height = viewport.height;
                            
                            // Clear canvas
                            context.clearRect(0, 0, pdfCanvas.width, pdfCanvas.height);
                            
                            // Render PDF page
                            pdfRenderTask = pdfPage.render({
                                canvasContext: context,
                                viewport: viewport
                            });
                            
                            await pdfRenderTask.promise;
                            pdfRenderTask = null;
                            
                            // If we're on page 2 and highlighting, draw the highlight
                            if (pageNum === 2 && window.highlightRequested) {
                                drawHighlight(context, viewport);
                                window.highlightRequested = false;
                            }
                            
                        } catch (error) {
                            if (error.name === 'RenderingCancelled') {
                                return; // Rendering was cancelled, do nothing
                            }
                            console.error('Error rendering page:', error);
                        }
                    }
                    
                    // Draw highlight rectangle
                    function drawHighlight(context, viewport) {
                        // Get word coordinates (scaled for viewport)
                        const x0 = """ + str(x0) + """ * scale;
                        const y0 = """ + str(y0) + """ * scale;
                        const x1 = """ + str(x1) + """ * scale;
                        const y1 = """ + str(y1) + """ * scale;
                        
                        // Since PDF coordinates start from bottom-left, we need to adjust y-coordinates
                        const top = viewport.height - y1;
                        const height = y1 - y0;
                        
                        // Draw highlight
                        context.save();
                        context.fillStyle = 'rgba(255, 255, 0, 0.5)';
                        context.strokeStyle = 'rgba(255, 0, 0, 0.8)';
                        context.lineWidth = 2;
                        context.fillRect(x0, top, x1 - x0, height);
                        context.strokeRect(x0, top, x1 - x0, height);
                        context.restore();
                        
                        // Scroll to highlight
                        scrollToHighlight(x0, top, x1 - x0, height);
                    }
                    
                    // Scroll to center the highlighted word
                    function scrollToHighlight(x, y, width, height) {
                        // Get scroll container (PDF section)
                        const container = document.querySelector('.pdf-section');
                        
                        // Calculate center of highlight
                        const highlightCenterX = x + (width / 2);
                        const highlightCenterY = y + (height / 2);
                        
                        // Get canvas position
                        const canvas = pdfCanvas;
                        const canvasRect = canvas.getBoundingClientRect();
                        
                        // Calculate scroll position to center highlight
                        const scrollX = highlightCenterX - (container.clientWidth / 2);
                        const scrollY = highlightCenterY - (container.clientHeight / 2);
                        
                        // Scroll to position
                        container.scrollTo({
                            left: scrollX,
                            top: scrollY,
                            behavior: 'smooth'
                        });
                    }
                    
                    // Go to PDF word function
                    function goToPdfWord() {
                        // Flag to indicate highlighting is needed
                        window.highlightRequested = true;
                        
                        // Navigate to page 2 and highlight
                        if (currentPageNum === 2) {
                            // If already on page 2, just highlight
                            const viewport = pdfPage.getViewport({ scale });
                            drawHighlight(pdfCanvas.getContext('2d'), viewport);
                        } else {
                            // Navigate to page 2
                            renderPage(2);
                        }
                    }
                    
                    // Initialize the PDF viewer
                    initPdfViewer();
                </script>
            </body>
            </html>
            """
            
            # Write to file
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            # Open in browser
            webbrowser.open('file://' + html_path)
            self.status_var.set("Linked document viewer created and opened in browser")
            
        except Exception as e:
            messagebox.showerror("Error", f"Error creating linked viewer: {str(e)}")


if __name__ == "__main__":
    root = tk.Tk()
    app = DocumentLinker(root)
    root.mainloop()